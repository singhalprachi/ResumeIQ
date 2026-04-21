import pdfplumber
import docx
import io
import re
from typing import Tuple
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from PDF preserving structure. Returns (text, page_count)."""
    text_parts = []
    page_count = 0
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text(layout=True)
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise ValueError(f"Could not parse PDF: {str(e)}")

    full_text = "\n\n".join(text_parts)
    return clean_text(full_text), page_count


def extract_text_from_docx(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from DOCX preserving paragraph structure. Returns (text, estimated_page_count)."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise ValueError(f"Could not parse DOCX: {str(e)}")

    full_text = "\n".join(paragraphs)
    
    # Estimate page count for DOCX based on word count (approx 400-500 words per page)
    word_count = len(full_text.split())
    estimated_page_count = max(1, round(word_count / 450))
    
    return clean_text(full_text), estimated_page_count


def clean_text(text: str) -> str:
    """Clean extracted text."""
    # Remove excessive whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    # Remove non-printable characters except newlines and tabs
    text = re.sub(r"[^\x20-\x7E\n\t]", " ", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """Split text into overlapping chunks for RAG."""
    words = text.split()
    chunks = []
    start = 0

    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        if end >= len(words):
            break
        start = end - overlap  # overlap for context continuity

    return [c for c in chunks if len(c.strip()) > 50]


def detect_resume_sections(text: str) -> dict[str, str]:
    """
    Attempt to identify major resume sections.
    Returns a dict of {section_name: section_text}.
    """
    section_patterns = {
        "contact": r"(?i)(contact|personal\s+info|email|phone)",
        "summary": r"(?i)(summary|objective|profile|about\s+me|professional\s+summary)",
        "experience": r"(?i)(experience|work\s+history|employment|career)",
        "education": r"(?i)(education|academic|qualification|degree|university|college)",
        "skills": r"(?i)(skills|technical\s+skills|competencies|technologies|tools)",
        "certifications": r"(?i)(certification|certificate|license|accreditation)",
        "projects": r"(?i)(projects|portfolio|work\s+samples)",
        "achievements": r"(?i)(achievement|award|honor|recognition)",
    }

    lines = text.split("\n")
    sections: dict[str, list[str]] = {}
    current_section = "header"
    sections[current_section] = []

    for line in lines:
        matched_section = None
        for section_name, pattern in section_patterns.items():
            if re.search(pattern, line) and len(line.strip()) < 60:
                matched_section = section_name
                break

        if matched_section:
            current_section = matched_section
            if current_section not in sections:
                sections[current_section] = []
        else:
            sections[current_section].append(line)

    return {k: "\n".join(v).strip() for k, v in sections.items() if v}


def parse_document(file_bytes: bytes, content_type: str) -> Tuple[str, dict[str, str], int]:
    """
    Main entry point. Returns (full_text, sections_dict, page_count).
    """
    if "pdf" in content_type:
        full_text, page_count = extract_text_from_pdf(file_bytes)
    elif "wordprocessingml" in content_type or "docx" in content_type:
        full_text, page_count = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {content_type}")

    if not full_text or len(full_text) < 100:
        raise ValueError("Could not extract meaningful text from document. Please check the file.")

    sections = detect_resume_sections(full_text)
    return full_text, sections, page_count
